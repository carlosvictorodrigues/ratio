from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from backend.escritorio.models import RatioEscritorioState


class FormatadorPeticao:
    BLACK = RGBColor(0, 0, 0)
    SECTION_ORDER = (
        "enderecamento",
        "qualificacao",
        "fatos",
        "fundamentacao_juridica_responsabilidade",
        "fundamentacao_juridica_danos_materiais",
        "fundamentacao_juridica_danos_morais",
        "fundamentacao_juridica_perda_chance",
        "pedidos",
        "valor_da_causa",
    )
    SUBSECTION_TITLES = {
        "fundamentacao_juridica_responsabilidade": "Da Responsabilidade Civil",
        "fundamentacao_juridica_danos_materiais": "Dos Danos Materiais",
        "fundamentacao_juridica_danos_morais": "Dos Danos Morais",
        "fundamentacao_juridica_perda_chance": "Da Perda de uma Chance",
    }
    EMPHASIS_RE = re.compile(
        r"(in re ipsa|Tema\s+\d+(?:\.\d+)?|(?:REsp|RE|ARE|AREsp|AgInt|AgRg|EDcl|EAREsp|EREsp)\s+\d[\d\.\-\/]*|art\.?\s*\d+[A-Za-zº°\-]*(?:,\s*§\s*\d+[A-Za-zº°\-]*)?)",
        re.IGNORECASE,
    )
    QUOTE_BLOCK_RE = re.compile(r"^(.*?:)\s*[\"'“”](.{60,}?)[\"'“”]\s*(.*)$", re.DOTALL)

    def __init__(self, *, output_dir: str | Path):
        self.output_dir = Path(output_dir).expanduser()
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.doc = Document()
        self._configure_document()

    def _configure_document(self) -> None:
        section = self.doc.sections[0]
        section.left_margin = Cm(3)
        section.top_margin = Cm(3)
        section.right_margin = Cm(2)
        section.bottom_margin = Cm(2)

        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.color.rgb = self.BLACK
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Cm(2)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for level in range(1, 4):
            heading = self.doc.styles[f"Heading {level}"]
            heading.font.name = "Times New Roman"
            heading.font.bold = True
            heading.font.size = Pt(14 if level == 1 else 12)
            heading.font.color.rgb = self.BLACK
            heading.paragraph_format.first_line_indent = Cm(0)
            heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_cabecalho(self, tribunal: str = "", vara: str = "", processo: str = "") -> None:
        header = self.doc.sections[0].header
        paragraph = header.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if tribunal:
            run = paragraph.add_run(f"{tribunal}\n")
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.font.color.rgb = self.BLACK
            run.bold = True
        if vara:
            run = paragraph.add_run(f"{vara}\n")
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.color.rgb = self.BLACK
        if processo:
            run = paragraph.add_run(f"Processo: {processo}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.color.rgb = self.BLACK

    def add_enderecamento(self, texto_enderecamento: str):
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(texto_enderecamento or "").strip().upper())
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = self.BLACK
        self.doc.add_paragraph("")
        self.doc.add_paragraph("")
        return paragraph

    def add_secao(self, titulo: str, conteudo: str) -> None:
        self.doc.add_heading(titulo, level=1)
        self.doc.add_paragraph(conteudo)

    def add_heading_text(self, titulo: str, *, level: int = 1) -> None:
        self.doc.add_heading(titulo, level=level)

    def add_paragraph_with_emphasis(self, texto: str) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        text = str(texto or "").strip()
        last_end = 0
        for match in self.EMPHASIS_RE.finditer(text):
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            token = match.group(0)
            run = paragraph.add_run(token)
            if token.lower() == "in re ipsa":
                run.italic = True
            else:
                run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.color.rgb = self.BLACK
            last_end = match.end()
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def add_citacao_longa(self, texto: str):
        paragraph = self.doc.add_paragraph(str(texto or "").strip())
        paragraph.paragraph_format.left_indent = Cm(4)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.color.rgb = self.BLACK
            run.italic = True
        return paragraph

    def add_structured_content(self, conteudo: str) -> None:
        paragraphs = [part.strip() for part in re.split(r"\n\s*\n", str(conteudo or "")) if part.strip()]
        if not paragraphs:
            return
        for paragraph_text in paragraphs:
            quote_match = self.QUOTE_BLOCK_RE.match(paragraph_text)
            if quote_match:
                prefix, quote, suffix = quote_match.groups()
                if prefix.strip():
                    self.add_paragraph_with_emphasis(prefix.strip())
                self.add_citacao_longa(quote.strip())
                if suffix.strip():
                    self.add_paragraph_with_emphasis(suffix.strip())
                continue
            self.add_paragraph_with_emphasis(paragraph_text)

    def add_citacao(self, texto: str, *, verificada: bool = True):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(texto)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.color.rgb = self.BLACK
        if verificada:
            run.bold = True
            run.italic = True
        else:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            marker = paragraph.add_run(" [VERIFICAR]")
            marker.font.name = "Times New Roman"
            marker.font.size = Pt(12)
            marker.font.bold = True
            marker.font.color.rgb = RGBColor(255, 0, 0)
            marker.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return paragraph

    def add_numeracao_paginas(self) -> None:
        footer = self.doc.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._element.append(fld_begin)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        run._element.append(instr)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._element.append(fld_end)

    def _save_with_fallback(self, preferred_path: Path) -> Path:
        try:
            self.doc.save(preferred_path)
            return preferred_path
        except PermissionError:
            for index in range(1, 21):
                candidate = preferred_path.with_name(f"{preferred_path.stem}-{index}{preferred_path.suffix}")
                try:
                    self.doc.save(candidate)
                    return candidate
                except PermissionError:
                    continue
            raise

    def gerar(self, state: RatioEscritorioState, verificacoes: list[dict]) -> str:
        self.add_cabecalho(processo=state.caso_id)
        sections = dict(state.peca_sections or {})
        enderecamento = getattr(state, "enderecamento", "") or sections.get("enderecamento", "") or ""
        if enderecamento:
            self.add_enderecamento(enderecamento)
            sections.pop("enderecamento", None)

        ordered_keys = [key for key in self.SECTION_ORDER if key in sections] + [key for key in sections if key not in self.SECTION_ORDER]
        rendered_fundamentos = False
        fundamento_index = 1

        for titulo in ordered_keys:
            conteudo = sections.get(titulo, "")
            if titulo == "qualificacao":
                self.add_heading_text("QUALIFICACAO", level=1)
                self.add_structured_content(conteudo)
                continue
            if titulo == "fatos":
                self.add_heading_text("1. DOS FATOS", level=1)
                self.add_structured_content(conteudo)
                continue
            if titulo.startswith("fundamentacao_juridica_"):
                if not rendered_fundamentos:
                    self.add_heading_text("2. DOS FUNDAMENTOS JURIDICOS", level=1)
                    rendered_fundamentos = True
                subsection_title = self.SUBSECTION_TITLES.get(
                    titulo,
                    titulo.replace("fundamentacao_juridica_", "").replace("_", " ").title(),
                )
                self.add_heading_text(f"2.{fundamento_index}. {subsection_title}", level=2)
                fundamento_index += 1
                self.add_structured_content(conteudo)
                continue
            if titulo == "pedidos":
                self.add_heading_text("3. DOS PEDIDOS", level=1)
                self.add_structured_content(conteudo)
                continue
            if titulo == "valor_da_causa":
                self.add_heading_text("4. DO VALOR DA CAUSA", level=1)
                self.add_structured_content(conteudo)
                continue

            self.add_heading_text(titulo.upper().replace("_", " "), level=1)
            self.add_structured_content(conteudo)

        for verificacao in verificacoes:
            referencia = str(verificacao.get("referencia") or "").strip()
            if not referencia:
                continue
            verificada = str(verificacao.get("level") or "") in {"exact_match", "strong_match"}
            self.add_citacao(referencia, verificada=verificada)

        self.add_numeracao_paginas()
        output_path = self.output_dir / "peticao_final.docx"
        saved_path = self._save_with_fallback(output_path)
        return str(saved_path)
