from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from backend.escritorio.formatter import FormatadorPeticao
from backend.escritorio.models import RatioEscritorioState


def test_formatador_peticao_gera_docx_com_secoes_do_state(tmp_path: Path):
    state = RatioEscritorioState(
        caso_id="caso-1",
        tipo_peca="peticao_inicial",
        peca_sections={
            "dos_fatos": "Texto dos fatos.",
            "do_direito": "Texto do direito.",
        },
    )
    formatter = FormatadorPeticao(output_dir=tmp_path)

    path = formatter.gerar(state, verificacoes=[])

    assert Path(path).exists()
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "DOS FATOS" in texts
    assert "Texto dos fatos." in texts
    assert "DO DIREITO" in texts
    assert "Texto do direito." in texts


def test_formatador_peticao_destaca_citacao_nao_verificada(tmp_path: Path):
    formatter = FormatadorPeticao(output_dir=tmp_path)
    paragraph = formatter.add_citacao("REsp 1.234.567", verificada=False)

    marker_text = "".join(run.text for run in paragraph.runs)
    assert "[VERIFICAR]" in marker_text


def test_formatador_configura_estilo_normal_justificado_com_espacamento_e_cor(tmp_path: Path):
    formatter = FormatadorPeticao(output_dir=tmp_path)
    style = formatter.doc.styles["Normal"]

    assert style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert style.paragraph_format.space_after == Pt(6)
    assert style.font.color.rgb == formatter.BLACK


def test_formatador_add_enderecamento_centralizado_em_caixa_alta(tmp_path: Path):
    formatter = FormatadorPeticao(output_dir=tmp_path)
    paragraph = formatter.add_enderecamento("Excelentíssimo Senhor Doutor Juiz de Direito")

    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraph.runs[0].text == "EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO"
    assert paragraph.runs[0].bold is True


def test_formatador_add_citacao_longa_usa_bloco_recuado_e_fonte_10(tmp_path: Path):
    formatter = FormatadorPeticao(output_dir=tmp_path)
    paragraph = formatter.add_citacao_longa("Texto da ementa em bloco.")

    assert round(paragraph.paragraph_format.left_indent.cm, 1) == 4.0
    assert paragraph.paragraph_format.line_spacing == 1.0
    assert paragraph.paragraph_format.first_line_indent == Cm(0)
    assert paragraph.runs[0].font.size == Pt(10)


def test_formatador_add_numeracao_paginas_injeta_campo_page_no_rodape(tmp_path: Path):
    formatter = FormatadorPeticao(output_dir=tmp_path)
    formatter.add_numeracao_paginas()

    footer_xml = formatter.doc.sections[0].footer.paragraphs[0]._element.xml
    assert "PAGE" in footer_xml


def test_formatador_peticao_salva_em_nome_alternativo_quando_docx_esta_bloqueado(tmp_path: Path):
    state = RatioEscritorioState(
        caso_id="caso-1",
        tipo_peca="peticao_inicial",
        peca_sections={"dos_fatos": "Texto dos fatos."},
    )
    formatter = FormatadorPeticao(output_dir=tmp_path)
    original_save = formatter.doc.save
    calls = []

    def fake_save(path):  # noqa: ANN001
        calls.append(Path(path).name)
        if len(calls) == 1:
            raise PermissionError("arquivo bloqueado")
        return original_save(path)

    formatter.doc.save = fake_save

    path = formatter.gerar(state, verificacoes=[])

    assert Path(path).exists()
    assert Path(path).name != "peticao_final.docx"
    assert Path(path).name.startswith("peticao_final-")


def test_formatador_peticao_cria_hierarquia_juridica_numerada(tmp_path: Path):
    state = RatioEscritorioState(
        caso_id="caso-1",
        tipo_peca="peticao_inicial",
        peca_sections={
            "qualificacao": "Autor qualificado.",
            "fatos": "Texto dos fatos.",
            "fundamentacao_juridica_responsabilidade": "Texto da responsabilidade.",
            "fundamentacao_juridica_danos_materiais": "Texto dos danos materiais.",
            "pedidos": "Texto dos pedidos.",
            "valor_da_causa": "Texto do valor.",
        },
    )
    formatter = FormatadorPeticao(output_dir=tmp_path)

    path = formatter.gerar(state, verificacoes=[])

    doc = Document(path)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "QUALIFICACAO" in texts
    assert "1. DOS FATOS" in texts
    assert "2. DOS FUNDAMENTOS JURIDICOS" in texts
    assert any(text.startswith("2.1. Da Responsabilidade") for text in texts)
    assert any(text.startswith("2.2. Dos Danos Materiais") for text in texts)
    assert "3. DOS PEDIDOS" in texts
    assert "4. DO VALOR DA CAUSA" in texts


def test_formatador_peticao_converte_tese_firmada_em_bloco_recuado(tmp_path: Path):
    state = RatioEscritorioState(
        caso_id="caso-1",
        tipo_peca="peticao_inicial",
        peca_sections={
            "fundamentacao_juridica_responsabilidade": (
                "A tese firmada estabelece que: "
                "'O Estado responde subsidiariamente por danos materiais causados aos candidatos.' "
                "Resta cristalino o dever de indenizar."
            )
        },
    )
    formatter = FormatadorPeticao(output_dir=tmp_path)

    path = formatter.gerar(state, verificacoes=[])

    doc = Document(path)
    quote_paragraph = next(
        p for p in doc.paragraphs
        if "O Estado responde subsidiariamente por danos materiais" in p.text
    )
    assert round((quote_paragraph.paragraph_format.left_indent.cm if quote_paragraph.paragraph_format.left_indent else 0), 1) == 4.0
    assert any(run.italic for run in quote_paragraph.runs if run.text.strip())


def test_formatador_peticao_destaca_referencias_e_latinismos_com_enfase(tmp_path: Path):
    state = RatioEscritorioState(
        caso_id="caso-1",
        tipo_peca="peticao_inicial",
        peca_sections={
            "fundamentacao_juridica_danos_morais": "O dano moral 'in re ipsa' foi reconhecido no RE 662.405."
        },
    )
    formatter = FormatadorPeticao(output_dir=tmp_path)

    path = formatter.gerar(state, verificacoes=[])

    doc = Document(path)
    paragraph = next(
        p for p in doc.paragraphs
        if "O dano moral" in p.text
    )
    assert any(run.italic for run in paragraph.runs if "in re ipsa" in run.text)
    assert any(run.bold for run in paragraph.runs if "RE 662.405" in run.text)
